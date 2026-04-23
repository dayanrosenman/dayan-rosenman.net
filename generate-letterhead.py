"""
Generates letterhead.docx — upload to Google Drive and open as a Google Doc.
The letterhead lives in the document header (repeats on every page).
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MAROON      = RGBColor(0x67, 0x0d, 0x00)
MAROON_DARK = RGBColor(0x3d, 0x08, 0x00)
INK_LIGHT   = RGBColor(0x8a, 0x7a, 0x72)
INK         = RGBColor(0x3a, 0x32, 0x30)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


def set_font(run, name="Georgia", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_bottom_border(paragraph, color="670D00", size=12):
    """Add a bottom border (rule) to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_top_border_to_section(section):
    """Add a maroon top page border to the section (appears at very top of page)."""
    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "24")   # 3pt line
    top.set(qn("w:space"), "24")
    top.set(qn("w:color"), "670D00")
    pgBorders.append(top)
    sectPr.append(pgBorders)


doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(0.5)
section.bottom_margin = Inches(0.75)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)
section.different_first_page_header_footer = False

add_top_border_to_section(section)

# ── Header ───────────────────────────────────────────────────────────────────
header = section.header

# Name
p_name = header.paragraphs[0]
p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_spacing(p_name, before=80, after=40)
run = p_name.add_run("David Dayan-Rosenman, MD")
set_font(run, name="Georgia", size=20, bold=True, color=MAROON_DARK)

# Credentials line 1
p_cred1 = header.add_paragraph()
set_paragraph_spacing(p_cred1, before=0, after=20)
run = p_cred1.add_run("Board Certified in Addiction Medicine & Internal Medicine")
set_font(run, name="Calibri", size=10, color=INK_LIGHT)

# Credentials line 2
p_cred2 = header.add_paragraph()
set_paragraph_spacing(p_cred2, before=0, after=60)
run = p_cred2.add_run("Harvard Medical School Faculty")
set_font(run, name="Calibri", size=10, color=INK_LIGHT)

# Specialty / location — with bottom border acting as the divider
p_specialty = header.add_paragraph()
p_specialty.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_spacing(p_specialty, before=0, after=80)
add_bottom_border(p_specialty, color="670D00", size=8)
run = p_specialty.add_run("MENTAL HEALTH & ADDICTION CARE  ·  CAMBRIDGE, MA")
set_font(run, name="Calibri", size=9, color=INK_LIGHT)
run.font.all_caps = True

# Contact bar
p_contact = header.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_spacing(p_contact, before=60, after=0)

def add_contact_item(para, label, value, add_separator=True):
    r_label = para.add_run(label.upper() + "  ")
    set_font(r_label, name="Calibri", size=8.5, color=INK_LIGHT)
    r_label.font.all_caps = True
    r_val = para.add_run(value)
    set_font(r_val, name="Calibri", size=9.5, color=INK)
    if add_separator:
        r_sep = para.add_run("    ·    ")
        set_font(r_sep, name="Calibri", size=9.5, color=INK_LIGHT)

add_contact_item(p_contact, "Phone", "(617) 528-0538")
add_contact_item(p_contact, "Email", "david@dayan-rosenman.net")
add_contact_item(p_contact, "Office", "1 Arnold Circle, Cambridge, MA 02139")
add_contact_item(p_contact, "Web", "dayan-rosenman.net", add_separator=False)

# ── Footer ───────────────────────────────────────────────────────────────────
footer = section.footer
p_footer = footer.paragraphs[0]
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p_footer, before=0, after=0)
run = p_footer.add_run(
    "David Dayan-Rosenman, MD  ·  1 Arnold Circle, Cambridge, MA 02139  "
    "·  (617) 528-0538  ·  dayan-rosenman.net"
)
set_font(run, name="Calibri", size=8, color=INK_LIGHT)

# ── Body placeholder ─────────────────────────────────────────────────────────
p_body = doc.add_paragraph()
set_paragraph_spacing(p_body, before=0, after=0)

# ── Save ─────────────────────────────────────────────────────────────────────
out = "letterhead.docx"
doc.save(out)
print(f"Saved: {out}")
