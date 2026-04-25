"""Generate letterhead.docx matching the HTML letterhead design."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ACCENT      = RGBColor(0x67, 0x0d, 0x00)
ACCENT_DARK = RGBColor(0x3d, 0x08, 0x00)
INK_LIGHT   = RGBColor(0x8a, 0x7a, 0x72)
CREAM_BG    = RGBColor(0xf9, 0xf5, 0xf2)

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(0)
section.bottom_margin = Inches(0)
section.left_margin   = Inches(0)
section.right_margin  = Inches(0)

# ── Helper: set paragraph shading ───────────────────────────────────────────
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ── Helper: paragraph spacing ───────────────────────────────────────────────
def set_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:spacing')
    pBdr.set(qn('w:before'), str(before))
    pBdr.set(qn('w:after'),  str(after))
    if line:
        pBdr.set(qn('w:line'), str(line))
        pBdr.set(qn('w:lineRule'), 'exact')
    pPr.append(pBdr)

# ── Helper: paragraph left/right indent ─────────────────────────────────────
def set_indent(para, left_twips, right_twips=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),  str(left_twips))
    ind.set(qn('w:right'), str(right_twips))
    pPr.append(ind)

# ── Helper: top border on paragraph ─────────────────────────────────────────
def add_top_border(para, color_hex, size=40, space=0):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    str(size))
    top.set(qn('w:space'), str(space))
    top.set(qn('w:color'), color_hex)
    pBdr.append(top)
    pPr.append(pBdr)

# ── Helper: bottom border on paragraph ──────────────────────────────────────
def add_bottom_border(para, color_hex, size=8, space=0):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), str(space))
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

ONE_INCH = 1440  # twips

# ── HEADER BLOCK ─────────────────────────────────────────────────────────────
# Thick top border via first paragraph
p_top = doc.add_paragraph()
shade_paragraph(p_top, 'f9f5f2')
add_top_border(p_top, '670d00', size=40, space=0)
set_indent(p_top, ONE_INCH, ONE_INCH)
set_spacing(p_top, before=220, after=0)

# Name
p_name = doc.add_paragraph()
shade_paragraph(p_name, 'f9f5f2')
set_indent(p_name, ONE_INCH, ONE_INCH)
set_spacing(p_name, before=0, after=40)
run = p_name.add_run('David Dayan-Rosenman, MD')
run.font.name = 'Georgia'
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = ACCENT_DARK

# Credential line 1
p_cred1 = doc.add_paragraph()
shade_paragraph(p_cred1, 'f9f5f2')
set_indent(p_cred1, ONE_INCH, ONE_INCH)
set_spacing(p_cred1, before=0, after=20)
run = p_cred1.add_run('Board Certified in Addiction Medicine & Internal Medicine')
run.font.name = 'Georgia'
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = ACCENT

# Credential line 2
p_cred2 = doc.add_paragraph()
shade_paragraph(p_cred2, 'f9f5f2')
set_indent(p_cred2, ONE_INCH, ONE_INCH)
set_spacing(p_cred2, before=0, after=60)
run = p_cred2.add_run('Instructor in Psychiatry, Harvard Medical School')
run.font.name = 'Georgia'
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = ACCENT

# Specialty line
p_spec = doc.add_paragraph()
shade_paragraph(p_spec, 'f9f5f2')
set_indent(p_spec, ONE_INCH, ONE_INCH)
set_spacing(p_spec, before=0, after=160)
add_bottom_border(p_spec, '670d00', size=6, space=4)
run = p_spec.add_run('Mental Health & Addiction Care  ·  Cambridge, MA')
run.font.name = 'Calibri'
run.font.size = Pt(8)
run.font.bold = True
run.font.color.rgb = INK_LIGHT

# ── BODY SPACER ──────────────────────────────────────────────────────────────
for _ in range(22):
    p = doc.add_paragraph()
    set_indent(p, ONE_INCH, ONE_INCH)
    set_spacing(p, before=0, after=0, line=240)

# ── FOOTER ───────────────────────────────────────────────────────────────────
p_footer = doc.add_paragraph()
shade_paragraph(p_footer, 'f9f5f2')
add_top_border(p_footer, '670d00', size=6, space=4)
set_indent(p_footer, ONE_INCH, ONE_INCH)
set_spacing(p_footer, before=80, after=80)

footer_text = '\u260e  (617) 528-0538    \u2709  david@dayan-rosenman.net    \u25ce  1 Arnold Circle, Cambridge, MA 02139    \u2299  www.dayan-rosenman.net'
run = p_footer.add_run(footer_text)
run.font.name = 'Calibri'
run.font.size = Pt(8)
run.font.color.rgb = ACCENT

out = '/Users/david/ClaudeCode/dayan-rosenman-net/letterhead.docx'
doc.save(out)
print(f'Saved: {out}')
